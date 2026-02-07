"""Analyze the reference docx file in detail."""
from docx import Document
from docx.oxml.ns import qn

doc = Document(r'E:\jd_jones_rag_complete\jd_jones_rag_complete\jd_jones_rag\docs\66361.docx')

# Detailed analysis of Table 2 (main products table)
table = doc.tables[2]
print("=== PRODUCTS TABLE DETAILED ===")
print(f"Rows: {len(table.rows)}, Cols: {len(table.columns)}")

# Check header rows for merge info
print()
for r_idx in range(2):
    row = table.rows[r_idx]
    print(f"Row {r_idx}:")
    for c_idx, cell in enumerate(row.cells):
        tc = cell._tc
        tcPr = tc.find(qn("w:tcPr"))
        merge_info = ""
        if tcPr is not None:
            vMerge = tcPr.find(qn("w:vMerge"))
            hMerge = tcPr.find(qn("w:gridSpan"))
            if vMerge is not None:
                val = vMerge.get(qn("w:val"), "continue")
                merge_info += f" vMerge={val}"
            if hMerge is not None:
                val = hMerge.get(qn("w:val"))
                merge_info += f" gridSpan={val}"
        text = cell.text.strip()[:40]
        print(f'  Cell [{c_idx}]{merge_info}: "{text}"')

# Check the last row (totals)
print()
print("=== LAST ROW (TOTALS) ===")
last_row = table.rows[-1]
for c_idx, cell in enumerate(last_row.cells):
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    merge_info = ""
    if tcPr is not None:
        hMerge = tcPr.find(qn("w:gridSpan"))
        if hMerge is not None:
            val = hMerge.get(qn("w:val"))
            merge_info += f" gridSpan={val}"
    text = cell.text.strip()[:50]
    print(f'  Cell [{c_idx}]{merge_info}: "{text}"')

# All paragraphs with their exact positions  
print()
print("=== ALL PARAGRAPHS (including empty) ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text[:120] if para.text else "(empty)"
    style = para.style.name if para.style else "None"
    align = para.alignment
    bold = any(run.bold for run in para.runs)
    print(f"P{i:02d} [{style}] align={align} bold={bold}: {text}")

# Check terms section between tables
print()
print("=== CHECKING FOR CONTENT BETWEEN TABLES ===")
print(f"Table 0 ends, Table 1 starts")
print(f"Table 1 ends, Table 2 starts") 
print(f"Table 2 ends, then terms/footer")
